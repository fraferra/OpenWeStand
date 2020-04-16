import json
import boto3
import os
import googlemaps
import re
import pyqrcode
import urllib
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
import os.path

dirname = os.path.dirname("__file__")

client = boto3.client("sns")

ssm_client = boto3.client("ssm")

s3 = boto3.resource('s3')
s3_client = boto3.client('s3') 

bucket = "stay-home-emails"

configPath = os.path.join(dirname, "client_secret.json")

dynamodb = boto3.resource('dynamodb')
table = dynamodb.Table('StayHomeFinal')

with open(configPath) as f:
    config = json.load(f)

googleMapsClient = googlemaps.Client(key=config.get("google_maps_key"))
lastIndexPath = config.get("last_index_parameter")

food_cols = config.get("food_cols")
macro_dict = config.get("macro_dict")
food_dict = config.get("food_dict")

def store_order(order_dict):
    table.put_item(
           Item=order_dict
        )

def get_last_index():
    return int(ssm_client.get_parameter(Name=lastIndexPath)["Parameter"]["Value"])

def get_value(address_components, key):
    for component in address_components:
        types = component.get("types", [])
        if key in types:
            return component["short_name"]
    return

def update_last_index(newValue):
    ssm_client.put_parameter(Name=lastIndexPath,
                             Value=str(newValue),
                             Type='String',
                             Overwrite=True)
    return newValue

def create_qr_and_store(localfile, info, folder):
    qr = pyqrcode.create(info)
    

    file_name = "/tmp/{}.png".format(localfile)
    object_name = "{}/{}.png".format(folder, localfile)

    qr.png(file_name, scale=5)

    try:
        response = s3_client.upload_file(file_name,
                                         bucket,
                                         object_name,
                                         ExtraArgs={'ACL': 'public-read'})
    except ClientError as e:
        return False

    return ("https://{}.s3-us-west-2.amazonaws.com/{}".format(bucket, object_name), file_name)

def add_formatted_paragraph(string, document, fontSize=9):
    p = document.add_paragraph()
    run = p.add_run(string)
    font = run.font
    font.name = 'Calibri'
    font.size = Pt(fontSize)
    p.paragraph_format.line_spacing = 0.65
    #p.paragraph_format.left_indent = Inches(1.5)
    p = document.paragraphs[-1] 
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    

def create_doc(event, file_name, file_name_status):
    document = Document()
    p = add_formatted_paragraph('QR Address',document)
    img = document.add_picture(file_name, width=Inches(1.25))

    p = document.paragraphs[-1] 
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    full_name = event["Full Name"]
    if full_name == "NOT_FOUND":
        full_name = event["Contact First Name"]

    p = add_formatted_paragraph('{} | {}'.format(event["BagId"], event["FormattedZipCode"]),document)
    p = add_formatted_paragraph('{}'.format( full_name ),document)
    p = add_formatted_paragraph('{}'.format(event["FormattedAddresses"]),document)
    p = add_formatted_paragraph('{}'.format(event["Delivery Instructions/Gate Co"]),document)
    p = add_formatted_paragraph('{} | {}'.format(event["RegionMacro"], event["Area"]),document)
    p = add_formatted_paragraph('{}'.format(event["Phone"]),document)
    p = add_formatted_paragraph('{}'.format(event["FoodList"]),document)
    p = add_formatted_paragraph('PACKED: [ ]',document)
    p = add_formatted_paragraph('QR Status',document)
    img = document.add_picture(file_name_status, width=Inches(1.00))
    last_paragraph = document.paragraphs[-1] 
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = add_formatted_paragraph('''Please be advised: We take many sanitary precautions\nto assure we’re keeping you as safe as possible.\n
    We recommend you take precautions to sanitize\nthe items in your bag as you unpack.\n
    Don’t forget to wash your hands!!\n''',document)
    p = add_formatted_paragraph('''We are 100% donation driven.\n 
    Any support is appreciated! 😀\n
    www.StayHomeSD.com/DONATE\n
    (619) 800-3252''',document)
    document.add_page_break()
    file_name_2 = '{}.docx'.format(file_name)
    document.save(file_name_2)
    return file_name_2

def create_label_and_store(event, file_name, file_name_status):
    new_file_name = create_doc(event, file_name, file_name_status)

    day = event["SubmissionTime"].split()[0]
    object_name_all = "labels/{}/all/{}_{}_{}_{}.docx".format(day, event["SubmissionTime"], event["RegionMacro"], event["Area"], event["FormattedAddresses"])
    object_name_zip = "labels/{}/{}/{}_{}_{}_{}.docx".format(day, event["FormattedZipCode"], event["SubmissionTime"], event["RegionMacro"], event["Area"], event["FormattedAddresses"])
    object_name_region = "labels/{}/{}/{}_{}_{}_{}_{}.docx".format(day, event["RegionMacro"], event["BagId"], event["SubmissionTime"], event["RegionMacro"], event["Area"], event["FormattedAddresses"])


    try:
        response = s3_client.upload_file(new_file_name,
                                         bucket,
                                         object_name_all,
                                         ExtraArgs={'ACL': 'public-read'})

        response = s3_client.upload_file(new_file_name,
                                         bucket,
                                         object_name_zip,
                                         ExtraArgs={'ACL': 'public-read'})

        response = s3_client.upload_file(new_file_name,
                                         bucket,
                                         object_name_region,
                                         ExtraArgs={'ACL': 'public-read'})
    except ClientError as e:
        return False

    return "https://{}.s3-us-west-2.amazonaws.com/{}".format(bucket, object_name_all)

def get_formatted_address(address, gmaps):
    try:
        geocode_results = gmaps.geocode(address)
        if len(geocode_results) == 1:
            geocode_result = geocode_results[0]
        elif len(geocode_results) > 1:
            for geo_result in geocode_results:
                key = "administrative_area_level_2"
                administrative_area_level_2 = get_value(geo_result["address_components"], key)
                if administrative_area_level_2 == 'San Diego County':
                    geocode_result = geo_result
                    break
        else:
            return
        zipCode = get_value(geocode_result["address_components"], "postal_code")
        formattedAddress = geocode_result["formatted_address"]
        neighborhood = get_value(geocode_result["address_components"], "neighborhood")
        geo_location = geocode_result.get("geometry", {}).get("location", {})
        if neighborhood is None:
            neighborhood = get_value(geocode_result["address_components"], "locality")
        return (zipCode, formattedAddress, neighborhood, geo_location)

    except Exception as e:
        return(address, address, address, address)


def shorten_neigh(neigh):
    try:
        return " ".join(neigh.split()[:2]).lower()
    except:
        return neigh

def augment_item(data):
    zipCode, formattedAddress, neighborhood, geo_location = get_formatted_address(data["Delivery Address"], googleMapsClient)
    data["FormattedAddresses"] = formattedAddress
    data["GeoLocation"] = str(geo_location)
    data["FormattedZipCode"] = zipCode
    data["Area"] = shorten_neigh(neighborhood)
    data["FoodList"] = ",".join(get_formatted_food_list(data))
    data["RegionMacro"] = macro_dict.get(shorten_neigh(neighborhood), "NOT_MATCHED")

    t = urllib.parse.quote_plus(data["FormattedAddresses"])
    addressInfo = "https://www.google.com/maps/place/{}".format(t)

    data["QrURL"], file_name = create_qr_and_store(t, addressInfo, "qrcodes")

    lastBagId = get_last_index()
    data["BagId"] = lastBagId
    update_last_index(lastBagId+1)
    data["SubmissionTime"] = data["Submission Time"]
    statusInfo = "https://skselrgwkk.execute-api.us-west-2.amazonaws.com/prod/status?bagId={}&token=stayhome&submissiontime={}".format(data["BagId"], data["SubmissionTime"])
    local_status = "bagId={}+submissiontime={}".format(data["BagId"], data["SubmissionTime"])
    data["StatusUrl"], file_name_status = create_qr_and_store(local_status, statusInfo, "qrstatus")

    data["LabelUrl"] = create_label_and_store(data, file_name, file_name_status)
    data["StatusUpdate"] = "NOT_DELIVERED"
 
    return data

def get_formatted_food_list(data):
        food_list = []
        for key, food_bool in data.items():
            if key in food_cols:
                if "Che" in food_bool:
                    food_list.append(food_dict[key])
        return sorted(food_list)

def lambda_handler(event, context):
    try:
        augmentItem = augment_item(event)
        store_order(augmentItem)
                    
        return {
            'statusCode': 200,
            'body': json.dumps('')
        }
    except Exception as e:
        print(e)
        return {
            'statusCode': 500,
            'body': json.dumps("")
        }

